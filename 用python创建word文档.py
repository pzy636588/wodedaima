from docx import Document
#创建文档
document=Document()
#标题
document.add_heading('标题0',0)
document.add_heading('标题1',1)
document.add_heading('标题2',2)
#保存
document.save('test.docx')